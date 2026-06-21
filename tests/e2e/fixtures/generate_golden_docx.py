#!/usr/bin/env python3
"""Bake the golden docx fixture used by the e2e suite.

Run once on a machine with python-docx installed (e.g. the aspirant-commander
venv); the produced file is checked in. Re-run only if the operator BR-flow
output contract changes shape.

The fixture mirrors what commander emits after the #886 fix:
- Body runs (`w:r/w:rPr/w:highlight`) are stripped — zero occurrences.
- Paragraph-mark formatting (`w:pPr/w:rPr/w:highlight`) is allowed — these
  residues are operator-invisible in the LibreOffice PDF render, per the
  #886 verification on PR #87.

Usage:
    cd ~/git/aspirant/aspirant-commander
    .venv/bin/python ~/git/aspirant/aspirant-client/tests/e2e/fixtures/generate_golden_docx.py
"""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def add_pmark_yellow_highlight(paragraph) -> None:
    """Attach <w:pPr><w:rPr><w:highlight w:val="yellow"/></w:rPr></w:pPr>.

    Mirrors the residue commander legitimately leaves behind after stripping
    body-run highlights — the structural-equivalence test must permit these.
    """
    p_elem = paragraph._p
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(p_elem, qn("w:pPr"))
        p_elem.insert(0, pPr)
    rPr = etree.SubElement(pPr, qn("w:rPr"))
    highlight = etree.SubElement(rPr, qn("w:highlight"))
    highlight.set(qn("w:val"), "yellow")


def main() -> None:
    out = Path(__file__).resolve().parent / "golden" / "vardeutlatande.golden.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Värdeutlåtande — BR Brf Exempel", level=1)

    p = doc.add_paragraph()
    p.add_run("Objekt: ").bold = True
    p.add_run("LGH 1001 Brf Exempel (769600-0000)")

    p = doc.add_paragraph()
    p.add_run("Adress: ").bold = True
    p.add_run("Exempelgatan 1, 123 45 Nynäshamn")

    p = doc.add_paragraph()
    p.add_run("Marknadsvärde: ").bold = True
    p.add_run("3 050 000 kr (± 50 000 kr)")

    pmark = doc.add_paragraph(
        "Nynäshamn, 2026-06-21 — Jenny Wiklund, Auktoriserad mäklare"
    )
    add_pmark_yellow_highlight(pmark)

    doc.save(out)
    print(f"wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
